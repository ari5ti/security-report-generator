﻿from docx.shared import Mm
from docx import Document
from docxcompose.composer import Composer
from docx import Document as Document_compose
from docxtpl import DocxTemplate, InlineImage, RichText
import os
import datetime, time
import csv
import random
import json

date = datetime.datetime.now().strftime('%d') + "-" + datetime.datetime.now().strftime('%b') + "-" + datetime.datetime.now().strftime('%Y')
mtemp = "mtemp.docx"
liste = []
vulns = []

# sanitize csv escape special characters
def sanitize_csv():
    ptkb = open("ptkb.csv", "r")
    ptkb = ''.join([i for i in ptkb]).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    x = open("ptkbs.csv", "w")
    x.writelines(ptkb)
    x.close()

# Function to convert a CSV to JSON
def make_json(csvfilepath, jsonfilepath):
    data = {}   # create a dictionary
    with open(csvfilepath, encoding='utf-8') as csvf:   # Open a csv reader called DictReader
        csvReader = csv.DictReader(csvf)
        
        for rows in csvReader:      # Convert each row into a dictionary and add it to data
            key = rows['id']      # Assuming a column named 'sl #' to be the primary key
            data[key] = rows
            if csvfilepath == "vuln.csv":
                rows['cap'] = list(rows['cap'].split('|'))
                rows['poc'] = list(rows['poc'].split('|'))

    with open(jsonfilepath, 'w', encoding='utf-8') as jsonf:
        jsonf.write(json.dumps(data, indent=4))       

# Function to replace client name & date in footer and client and project name in header
def head_foot(client, service):
    document = Document('temp.docx')
    section = document.sections[0]
    header = section.header
    footer = section.footer
    header.paragraphs[0].text = header.paragraphs[0].text.replace("Clientservice", client + ' - ' + service)
    footer.paragraphs[0].text = footer.paragraphs[0].text.replace("Client", client)
    footer.paragraphs[0].text = footer.paragraphs[0].text.replace("Date", date)
    document.save('vgen.docx')

# take user input regarding details to be filled
def user_input():
    global client, service, vendor, vpn, service_type, service_scope, service_detailed_scope, start_date, end_date, author1, author2
    client = input("Enter Client name: ")
    vendor = input("Enter Vendor name: ")
    service = input("Enter Service name(n/w vapt, web vapt): ")
    service_type = input("Enter Type of service(black, grey, white): ")
    service_scope = input("Enter Service scope(n/w, web app etc): ")
    service_detailed_scope = input("Enter Detailed Service scope(firewall, servers, wifi etc): ")
    start_date = input("Enter start date(July 1): ")
    end_date = input("Enter end date(July 5): ")
    author1 = input("Enter Author-1 name: ")
    author2 = input("Enter Author-2 name: ")
    vpn = input("Enter the VPN used: ")
    print("\nPlease wait, report is being generated.")
    print("\nPlease do not touch any files while report is being generated.")

# change risk color 
def vuln_color(n):
    critical_color = "cc0500"
    high_color = "df3d03"
    medium_color = "f9a009"
    low_color = "ffcb0d"

    if kb[n]['Risk'] == "CRITICAL":
        bgcolor = critical_color
        fcolor = "#ffffff"
    elif kb[n]['Risk'] == "HIGH":
        bgcolor = high_color
        fcolor = "#ffffff"
    elif kb[n]['Risk'] == "MEDIUM":
        bgcolor = medium_color
        fcolor = "#000000"
    else:
        bgcolor = low_color
        fcolor = "#000000"
    
    return bgcolor, fcolor

# fill all the general stuff
def fill_general():
    template = DocxTemplate("master-template.docx")  # In same directory
    global ccritical, chigh, cmedium, clow
    ccritical = chigh = cmedium = clow = 0
    if author1 and author2:
        authora = author1 + " and " + author2
        authorc = author1 + ", " + author2
    elif not author1: authora = authorc = author2
    else: authora = authorc = author1
    table_contents = []
    for i in range(0, count):
        n = vulns[i]
        if kb[n]['Risk'] == "CRITICAL": ccritical += 1
        elif kb[n]['Risk'] == "HIGH": chigh += 1
        elif kb[n]['Risk'] == "MEDIUM": cmedium += 1
        else: clow += 1
        bgcolor, fcolor = vuln_color(n)
        risk = RichText()
        risk.add(kb[n]['Risk'], color=fcolor, size=20)
        table_contents.append({
            'title': kb[n]['Vulnerability'],
            'description': kb[n]['descrp'],
            'risk': risk,
            'bgcolor': bgcolor,
            'index': i
            })
    risk_count = str(ccritical) + " critical risk, " + str(chigh) + " high risk, " + str(cmedium) + " medium risk and " + str(clow) + " low risk"
    sev = []
    top_sev = []
    exec_top1 = []
    exec_top2 = []
    sev.extend([ccritical, chigh, cmedium, clow])
    scount = 0  
    for i in range(0,4):
        if (scount < 2) and sev[i] != 0:
            scount += 1
            top_sev.append(i)
    if len(top_sev) == 1:
        top1 = top_sev[0]
        top2 = None
    elif top_sev[1] != 3:
        top1 = top_sev[0]
        top2 = top_sev[1]
    else:
        top1 = top_sev[0]
        top2 = None
    sev_text = ["CRITICAL", "HIGH", "MEDIUM", "LOW"]
    if top1 is not None and top2 is not None:
        top1 = sev_text[top1]
        top2 = sev_text[top2]
        for i in range(0, count):
            n = vulns[i]
            if kb[n]['Risk'] == top1:  
                exec_top1.append({
                'title': kb[n]['Vulnerability'],
                'description': kb[n]['descrp'],
                })
            if kb[n]['Risk'] == top2:  
                exec_top2.append({
                'title': kb[n]['Vulnerability'],
                'description': kb[n]['descrp'],
                })
    else: 
        top1 = sev_text[top1]
        for i in range(0, count):
            n = vulns[i]
            if kb[n]['Risk'] == top1:  
                exec_top1.append({
                'title': kb[n]['Vulnerability'],
                'description': kb[n]['descrp'],
                })
    context = {
        'service': service,
        'vendor': vendor,
        'client': client,
        'authorc': authorc,
        'authora': authora,
        'author1': author1,
        'author2': author2,
        'service_type': service_type,
        'service_scope': service_scope,
        'service_detailed_scope': service_detailed_scope,
        'start_date': start_date,
        'end_date': end_date,
        'date': date,
        'table_contents': table_contents,
        'vpn': vpn,
        'risk_count': risk_count,
        'top1_title': top1,
        'top2_title': top2,
        'top1': top1,
        'top2': top2,
        'exec_top1': exec_top1,
        'exec_top2': exec_top2
        }
    template.render(context)
    template.save("mtemp.docx")

# join all the vulns into 1 file
def join_vulns(mtemp,files_list):
    count=len(files_list)
    master = Document_compose(mtemp)
    composer = Composer(master)
    for i in range(0, count):
        doc_temp = Document_compose(files_list[i])
        composer.append(doc_temp)
    composer.save("temp.docx")

# generate separate file for each vuln
def split_vuln(i, n):
    template = DocxTemplate("vuln-template.docx")  # In same directory
    bgcolor, fcolor = vuln_color(n)
    risk = RichText()
    risk.add(kb[n]['Risk'], color=fcolor, size=20)
    images = []
    img = dict(zip(si[n]['poc'], si[n]['cap']))
    for k in img:
        info = {
            "image": InlineImage(template, image_descriptor="poc/"+k+".png", width=Mm(150)),
            "desc": img[k],
        }
        images.append(info)
    context = {
       'title': kb[n]['Vulnerability'],
       'risk': risk,
       'impact': kb[n]['Impact'],
       'likelihood': kb[n]['Likelihood'],
       'description': kb[n]['descrp'],
       'imp': kb[n]['imp'],
       'recommendation': kb[n]['recommendation'],
       'info': kb[n]['info'],
       'image': images,
       'bgcolor': bgcolor 
        }
    template.render(context)
    template.save("%s.docx" % str(i + 1))
    i = str(i + 1) + ".docx"
    liste.append(i)

    # Wait a random time - increase to (60,180) for real production run.
    wait = time.sleep(random.randint(1, 2))

# check all required files exists
def check_files():
    files = ['master-template.docx', 'vuln-template.docx', 'vuln.csv', 'ptkb.csv', 'poc']
    stat = True
    for f in files:
        if not os.path.exists(f):
            print(f, "file not found")
            stat = False
    if not stat:
        print("check all required files are in the same directory.")
        exit()

# clean the system
def cleanup():
    liste.extend(['ptkb.json', 'vuln.json', 'mtemp.docx', 'temp.docx', 'ptkbs.csv'])
    for f in liste:
        os.remove(f)

# info for the user
def info():
    print("\nReport Generated Successfully.\n")
    print("Open vgen.docx in windows, accept all errors/risk and save the file with required name.")
    print("refer the comments in the document to make necessary changes which were not able to achieve through automation.")
    
# generate word report
def generate_word_report():
    check_files()
    sanitize_csv()
    make_json("ptkbs.csv", "ptkb.json")
    make_json("vuln.csv", "vuln.json")
    with open("ptkb.json") as json_file:
        global kb, si, count
        kb = json.load(json_file)
    with open("vuln.json") as json_file:
        si = json.load(json_file)
    for k in si.keys():
        if k != "": vulns.append(k)
    count = len(vulns)
    user_input()
    fill_general()
    for i in range(0, count):
        n = vulns[i]
        split_vuln(i, n)
    join_vulns(mtemp, liste)
    head_foot(client, service)
    cleanup()
    info()


generate_word_report()

